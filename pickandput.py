import docx
import pandas as pd
import win32com.client as client

def commify(n):
    """
    Add commas to an integer `n`.

        >>> commify(1)
        '1'
        >>> commify(123)
        '123'
        >>> commify(1234)
        '1,234'
        >>> commify(1234567890)
        '1,234,567,890'
        >>> commify(123.0)
        '123.0'
        >>> commify(1234.5)
        '1,234.5'
        >>> commify(1234.56789)
        '1,234.56789'
        >>> commify('%.2f' % 1234.5)
        '1,234.50'
        >>> commify(None)
        >>>

    """
    if n is None: return None
    n = str(n)
    if '.' in n:
        dollars, cents = n.split('.')
    else:
        dollars, cents = n, None

    r = []
    for i, c in enumerate(str(dollars)[::-1]):
        if i and (not (i % 3)):
            r.insert(0, ',')
        r.insert(0, c)
    out = ''.join(r)
    if cents:
        out += '.' + cents
    return out

'''
def convert_to_pdf(filepath: str):
    """Save a pdf of a docx file."""
    try:
        word = client.DispatchEx("Word.Application")
        target_path = filepath.replace("Updated Schedule Drivers & Vehicle.docx", r".pdf")
        word_doc = word.Documents.Open(filepath)
        word_doc.SaveAs("Updated Schedule Drivers & Vehicle", FileFormat=17)
        word_doc.Close()
    except Exception as e:
            raise e
    finally:
            word.Quit()


'''

#coverage = pd.read_excel("sampleform.xlsx", sheet_name="Sheet1")
#fleetSchedule = pd.read_excel("sampleform.xlsx", sheet_name="Sheet2")
#drivers = pd.read_excel("sampleform.xlsx", sheet_name="Sheet3")
flag = 0
count = 0
driverSchedule = pd.read_excel("Schedules Driver & Vehicle.xlsx", sheet_name="Sheet1")
equipmentSchedule = pd.read_excel("Schedules Driver & Vehicle.xlsx", sheet_name="Sheet2")
coverages = pd.read_excel("Schedules Driver & Vehicle.xlsx", sheet_name="Sheet3")
misc = pd.read_excel("Schedules Driver & Vehicle.xlsx", sheet_name="Sheet4")

labelsNew1 = driverSchedule.columns.values
labelsNew2 = equipmentSchedule.columns.values
labelsNew3 = coverages.columns.values

#labels = coverage.columns.values
#labels1 = fleetSchedule.columns.values
#labels2 = drivers.columns.values

#doc = docx.Document("NewTestDoc2.docx")
new = docx.Document("Schedules Driver & Vehicle.docx")

#buffer = coverage.at[0, labels[0]]
#print(buffer)
#print(labels)

#Sheet1
'''for i in range(coverage['coverage'].count()):
    paraDoc = doc.paragraphs[6+i]
    #paraDoc = doc.add_paragraph()

    for j in range(len(labels)):

        paraDoc.add_run("{}".format(coverage.at[i, labels[j]]))

        if j == 0:
            s = coverage.at[i, labels[0]]
            diff = 38 - len(s)
            if diff > 0:
                for k in range(diff):
                    paraDoc.add_run(" ")

        if j == 1:
            s = coverage.at[i, labels[1]]
            diff = 45 - len(s)
            if diff > 0:
                for k in range(diff):
                    paraDoc.add_run(" ")

        if j == 2:
            s = coverage.at[i, labels[2]]
            diff = 32 - len(s)
            if diff > 0:
                for k in range(diff):
                    paraDoc.add_run(" ")

        if j == 3:
            s1 = coverage.at[i, labels[3]]
            s = str(s1)
            diff = 17 - len(s)
            if diff > 0:
                for k in range(diff):
                    paraDoc.add_run(" ")

        #if j == 4:
            #s = coverage.at[i, labels[4]]
            #diff = 16 - len(s)
            #if diff > 0:
                #for k in range(diff):
                    #paraDoc.add_run(" ")'''
'''
#Sheet 1
#print(doc.tables[0].cell(0, 0).text)
for i in range(coverage['coverage'].count()):
    if i > 5:
        doc.tables[0].add_row()
    for j in range(len(labels)):
        buffer = str(coverage.at[i, labels[j]])
        doc.tables[0].cell(i+1, j).text = buffer



#Sheet 2
#print(doc.tables[1].cell(0, 0).text)
for i in range(fleetSchedule['year'].count()):
    if i > 3:
        doc.tables[1].add_row()
    for j in range(len(labels1)):
        buffer = str(fleetSchedule.at[i, labels1[j]])
        doc.tables[1].cell(i+1, j).text = buffer


#Sheet 3
#print(doc.tables[2].cell(0,0).text)
for i in range(drivers['Driver Name'].count()):
    if i > 1:
        doc.tables[2].add_row()
    for j in range(len(labels2)):
        buffer = str(drivers.at[i, labels2[j]])
        doc.tables[2].cell(i+1, j).text = buffer

#print(doc.paragraphs[9].runs[6].text)


#doc.tables[2].add_column(1)
#doc.tables[2].cell(0, 3).text = "New Column"



#doc.save("NewTestDoc.docx")

'''
#########################################################################################################################

#Sheet3 NewFile
for i in range(coverages[labelsNew3[0]].count()):
    #count += 1
    #print("count = {}".format(count))
    if i > 0:
        new.tables[0].add_row()
    for j in range(len(labelsNew3)):
        if j == 3 or j == 4:
            buffer = str(coverages.at[i, labelsNew3[j]])
            newBuffer = "$"+commify(buffer)
            new.tables[0].cell(i + 1, j).text = newBuffer
            continue

        buffer = str(coverages.at[i, labelsNew3[j]])
        new.tables[0].cell(i+1, j).text = buffer


#Sheet1_NewFile
for i in range(driverSchedule[labelsNew1[0]].count()):
    count += 1
    #print("count = {}".format(count))
    if i > 0:
        new1 = new.tables[1].add_row()
    for j in range(len(labelsNew1)):
        if j == 3 or j == 4:
            buffer = str(driverSchedule.at[i, labelsNew1[j]])
            newBuffer = buffer[5:10]
            newBuffer1 = newBuffer + "-" + buffer[0:4]
            new.tables[1].cell(i + 1, j).text = newBuffer1
            continue

        buffer = str(driverSchedule.at[i, labelsNew1[j]])
        new.tables[1].cell(i+1, j).text = buffer

total = 0
countEquip = 0

#Sheet2 NewFile
for i in range(equipmentSchedule[labelsNew2[0]].count()):
    count += 1
    countEquip += 1
    total += equipmentSchedule.at[i, labelsNew2[4]]
    #print("count = {}".format(count))
    if i > 0:
        new.tables[2].add_row()
    for j in range(len(labelsNew1)):
        if j == 4:
            buffer = str(commify(equipmentSchedule.at[i, labelsNew2[j]]))
            newBuffer = "$"+buffer
            new.tables[2].cell(i + 1, j).text = newBuffer
            break

        buffer = str(equipmentSchedule.at[i, labelsNew2[j]])
        new.tables[2].cell(i + 1, j).text = buffer

    if count > 19:
        flag = i
        #print("flag = {}".format(flag))
        break

#print("equipment = {}   flag ={}".format(equipmentSchedule[labelsNew1[0]].count(), flag))
total1 = total
str_total1 = str(commify(total1))
#new.paragraphs[4].text = "Total : $"+str_total1

#Table 3
if flag!=0:
    newFlag = equipmentSchedule[labelsNew2[0]].count() - countEquip
    for i in range(newFlag):
        total += equipmentSchedule.at[i, labelsNew2[4]]
        if i > 0:
            new.tables[3].add_row()
        for j in range(len(labelsNew2)):
            if j == 4:
                buffer = str(commify(equipmentSchedule.at[i, labelsNew2[j]]))
                newBuffer = "$" + buffer
                new.tables[3].cell(i + 1, j).text = newBuffer
                break

            buffer = str(equipmentSchedule.at[i+countEquip, labelsNew2[j]])
            new.tables[3].cell(i+1, j).text = buffer

total2 = total-total1
str_total2 = str(commify(total2))
str_total = str(commify(total))
#new.paragraphs[5].text = "Total : $"+str_total2
#new.paragraphs[6].text = "                                                                                                                                                       Total Insured Value : $"+str_total

buffer1 = str(misc.at[0, 'Effective Date'])
newBuffer = buffer1[5:10]
newBuffer1 = newBuffer+"-"+buffer1[0:4]

new.paragraphs[1].runs[0].text = " "+misc.at[0, 'Company']
new.paragraphs[2].runs[0].text = " "+misc.at[0, 'Address']
new.paragraphs[3].runs[3].text = " "+newBuffer1
new.paragraphs[6].runs[4].text = "$"+commify(str(misc.at[0, 'Down Payments']))+" and "+str(misc.at[0, 'Months'])+" monthly payments of $"+commify(str(misc.at[0, 'Installments']))
new.paragraphs[9].runs[2].text = "Towing and storage $"+commify(str(misc.at[0, 'Towing and Storage']))+". Graduated deductibles"
new.paragraphs[10].runs[2].text = "Earned freight $"+commify(str(misc.at[0, 'Earned Freight']))+", debris removal $"+commify(str(misc.at[0, 'Debris Removal']))+". Cargo Unattended Vehicle Exclusion. $"+commify(str(misc.at[0, 'Unattended Vehicle Exclusion']))+" sublimit for theft of target commodities. Reefer breakdown deductible $"+commify(str(misc.at[0, 'Breakdown deductable']))+" for units 10+ years of age. Reefer units must be maintained as per manufacturer’s guidelines every 30 days and record be kept and provided in case of a loss. See policy for details, terms, conditions, limitations, and exclusions."
new.paragraphs[45].runs[0].text = misc.at[0, 'Company']
new.paragraphs[46].runs[0].text = misc.at[0, 'Address']
new.paragraphs[47].runs[1].text = " "+newBuffer1
new.paragraphs[25].text = "Total: $"+str_total1
new.paragraphs[27].text = "Total: $"+str_total2
new.paragraphs[28].text = "                                                                                                                                                       Total Insured Value : $"+str_total
new.paragraphs[175].runs[0].text = "Your policy has an Auto Liability limit of $"+commify(str(misc.at[0,'Liability Limit']))+" only. This is the maximum amount your insurance company will pay in an accident where you are liable for damages."
#print(new.paragraphs[25].text);
#new1.style="ColorfulShading"
new.save("Updated_Schedule_Drivers_&_Vehicles.docx")
#print(new.paragraphs[24].runs[1].text)


